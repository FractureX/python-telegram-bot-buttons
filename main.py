from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox
from PIL import Image, ImageTk
from telebot import TeleBot, types
import threading
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Datos de ejemplo
from static.Data import id_jefe, personal, empresas, bancos, formas_pago, imagenes

# Inicializa el bot con tu token
bot = TeleBot("7537043769:AAHl4u70EoZVPuLzBuQNV_ZIulyTgwPALX0")  # Reemplaza con tu API KEY de Telegram

def validar_numero(valor):
    return valor == "" or valor.isdigit() or (valor.count('.') == 1 and valor.replace('.', '').isdigit())

def validar_alfanumerico(valor):
    return valor == "" or valor.isalnum()

def actualizar_bancos(event):
    empresa_seleccionada = combo_empresa.get()
    imagen_empresa = imagenes.get(empresa_seleccionada)

    combo_banco['values'] = bancos.get(empresa_seleccionada)
    combo_banco.current(0)

    if imagen_empresa:
        try:
            nueva_imagen = Image.open(f"static/images/{imagen_empresa}")
            nueva_imagen = nueva_imagen.resize((200, 100))
            img_tk = ImageTk.PhotoImage(nueva_imagen)
            label_imagen.config(image=img_tk)
            label_imagen.image = img_tk
            label_imagen.grid(row=0, column=0, columnspan=2, pady=10, sticky=tk.N)
        except Exception as e:
            print(f"No se pudo cargar la imagen {imagen_empresa}: {e}")
    else:
        label_imagen.config(image='')

def set_entry_text(field, text):
    field.config(state="normal")
    field.delete(0, tk.END)
    field.insert(0, text)
    field.config(state="readonly")

def enviar_confirmacion():
    datos = validar_datos()
    if datos:
        markup = types.InlineKeyboardMarkup(row_width=2)
        aprobado = types.InlineKeyboardButton('Aprobado', callback_data='aprobado')
        denegado = types.InlineKeyboardButton('Denegado', callback_data='denegado')
        markup.add(aprobado, denegado)

        mensaje = f'Se requiere que confirme "{datos.get("descripcion")}", por un monto en BsS de {datos.get("monto_bs")}, solicitado por {datos.get("solicitado_por")}, con una forma de pago de {datos.get("forma_pago")}'
        bot.send_message(id_jefe, mensaje, reply_markup=markup)
        entry_estatus.config(state="normal")
        entry_estatus.delete(0, tk.END)
        entry_estatus.insert(0, "Notificación enviada, esperando respuesta")
        combo_solicitante.config(state="disabled")
        combo_empresa.config(state="disabled")
        entry_descripcion.config(state="disabled")
        combo_forma_pago.config(state="disabled")
        combo_banco.config(state="disabled")
        entry_monto_bs.config(state="readonly")
        entry_monto_dolar.config(state="readonly")
        entry_tasa_dia.config(state="readonly")
        entry_estatus.config(state="readonly")
        btn_confirmar.config(state="disabled")
        
        set_entry_text(entry_estatus, "Notificación enviada")

# Manejador para capturar las respuestas cuando se presionan los botones
@bot.callback_query_handler(func=lambda call: True)
def callback_query(call):
    if call.data == "aprobado":
        bot.answer_callback_query(call.id, "Has aprobado el pago.")
        bot.send_message(call.message.chat.id, "Has confirmado que el pago será realizado.")
        set_entry_text(entry_estatus, "Aprobado")
        
        generar_docx()
        bot.send_document(call.message.chat.id, types.InputFile("datos_pago.docx"))
        messagebox.showinfo("Éxito", "Los datos han sido generados y guardados en 'datos_pago.docx'.")
    elif call.data == "denegado":
        bot.answer_callback_query(call.id, "Has denegado el pago.")
        bot.send_message(call.message.chat.id, "Has denegado que el pago iba a realizarse.")
        set_entry_text(entry_estatus, "Denegado")

def iniciar_bot():
    bot.polling(timeout=3600)

def validar_datos() -> dict[str, any] | None:
    url_logo = f"static/images/{imagenes.get(combo_empresa.get())}"
    descripcion = entry_descripcion.get("1.0", tk.END).strip()
    forma_pago = combo_forma_pago.get()
    banco = combo_banco.get()
    monto_dolar = entry_monto_dolar.get()
    monto_bs = entry_monto_bs.get()
    tasa_dia = entry_tasa_dia.get()
    solicitado_por = combo_solicitante.get()
    estatus = entry_estatus.get()

    if not (descripcion and forma_pago and banco and solicitado_por and estatus):
        messagebox.showerror("Error", "Los campos alfabéticos no pueden estar vacíos.")
        return
    
    if not (validar_numero(monto_dolar) and validar_numero(monto_bs) and validar_numero(tasa_dia)):
        messagebox.showerror("Error", "Los campos numéricos deben ser válidos.")
        return
    
    return {
        "url_logo": url_logo,
        "descripcion": descripcion,
        "forma_pago": forma_pago,
        "banco": banco if banco else "N/A",
        "monto_dolar": monto_dolar if monto_dolar else "N/A",
        "monto_bs": monto_bs,
        "tasa_dia": tasa_dia if tasa_dia else "N/A",
        "solicitado_por": solicitado_por,
        "estatus": estatus
    }

def generar_docx():
    datos = validar_datos()

    doc = Document()
    doc.add_picture(image_path_or_stream=datos.get("url_logo"))
    
    doc.add_paragraph("")
    
    p_fecha = doc.add_paragraph()
    p_fecha.add_run('Fecha: ').bold = True
    p_fecha.add_run(datetime.strftime(datetime.now(), "%d/%m/%Y"))
    
    p_hora = doc.add_paragraph()
    p_hora.add_run('Hora: ').bold = True
    p_hora.add_run(datetime.strftime(datetime.now(), "%I:%M %p"))
    
    doc.add_paragraph("")
    
    solicitud = doc.add_paragraph("SOLICITUD DE APROBACIÓN")
    solicitud.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    solicitud.runs[0].bold = True
    solicitud.runs[0].underline = True
    
    table = doc.add_table(8, 2)
    table.style = "TableGrid"
    col_1_data = ["Descripción", "Forma de pago", "Banco", "Monto $", "Monto Bs", "Tasa del día", "Solicitado por", "Estatus"]
    col_2_data = [datos.get("descripcion"), datos.get("forma_pago"), datos.get("banco"), datos.get("monto_dolar"), datos.get("monto_bs"), datos.get("tasa_dia"), datos.get("solicitado_por"), datos.get("estatus")]
    
    # Añadir la parte izquierda
    for row_index in range(len(table.rows)):
        # Establecer el color de fondo para la primera columna
        cell = table.cell(row_index, 0)
        cell.width = Inches(1.28)
        
        shading_elm = parse_xml(r'<w:shd {} w:fill="2E74B5"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading_elm)  # Color de fondo
        cell.text = col_1_data[row_index]
        
        cell2 = table.cell(row_index, 1)
        cell2.width = cell2.width + (cell2.width - Inches(1.28))
        cell2.text = col_2_data[row_index]
        
        # Establecer el color de la fuente en la primera columna
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
        
        # Establecer el color de la fuente en la segunda columna
        for paragraph in cell2.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_picture(image_path_or_stream="static/images/firma.png")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("_________________")
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Pdte. José Goitia")
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Iterar sobre todos los párrafos del documento
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
    
    doc.save('datos_pago.docx')

# Crear la ventana principal
root = tk.Tk()
root.title("Formulario de Pago")
root.geometry("500x600")
root.resizable(False, False)

# Crear un frame para centrar el contenido
frame_contenido = ttk.Frame(root)
frame_contenido.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

# Configurar imagen
try:
    imagen = Image.open("static/images/logo.png")
    imagen = imagen.resize((200, 100))
    img_tk = ImageTk.PhotoImage(imagen)
    label_imagen = ttk.Label(frame_contenido, image=img_tk)
    label_imagen.grid(row=0, column=0, columnspan=2, pady=10, sticky=tk.EW)
except Exception as e:
    print(f"No se pudo cargar la imagen: {e}")
    label_imagen = ttk.Label(frame_contenido)
    label_imagen.grid(row=0, column=0, columnspan=2, pady=10, sticky=tk.EW)

# Validaciones de entrada
vc_numero = root.register(validar_numero)
vc_alfanumerico = root.register(validar_alfanumerico)

ttk.Label(frame_contenido, text="Solicitante:").grid(row=1, column=0, padx=5, pady=5, sticky=tk.E)
combo_solicitante = ttk.Combobox(frame_contenido, values=personal, state="readonly")
combo_solicitante.grid(row=1, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Empresa:").grid(row=2, column=0, padx=5, pady=5, sticky=tk.E)
combo_empresa = ttk.Combobox(frame_contenido, values=empresas, state="readonly")
combo_empresa.grid(row=2, column=1, padx=5, pady=5, sticky=tk.EW)
combo_empresa.bind("<<ComboboxSelected>>", actualizar_bancos)

# Crear etiquetas y campos de entrada
ttk.Label(frame_contenido, text="Descripción:").grid(row=3, column=0, padx=5, pady=5, sticky=tk.E)
entry_descripcion = tk.Text(frame_contenido, width=30, height=5)
entry_descripcion.grid(row=3, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Forma de pago:").grid(row=4, column=0, padx=5, pady=5, sticky=tk.E)
combo_forma_pago = ttk.Combobox(frame_contenido, values=formas_pago, state="readonly", width=30)
combo_forma_pago.grid(row=4, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Banco:").grid(row=5, column=0, padx=5, pady=5, sticky=tk.E)
combo_banco = ttk.Combobox(frame_contenido, width=30, state="readonly")
combo_banco.grid(row=5, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Monto $:").grid(row=6, column=0, padx=5, pady=5, sticky=tk.E)
entry_monto_dolar = ttk.Entry(frame_contenido, validate="key", validatecommand=(vc_numero, '%P'))
entry_monto_dolar.grid(row=6, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Monto BsS:").grid(row=7, column=0, padx=5, pady=5, sticky=tk.E)
entry_monto_bs = ttk.Entry(frame_contenido, validate="key", validatecommand=(vc_numero, '%P'))
entry_monto_bs.grid(row=7, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Tasa del día:").grid(row=8, column=0, padx=5, pady=5, sticky=tk.E)
entry_tasa_dia = ttk.Entry(frame_contenido, validate="key", validatecommand=(vc_numero, '%P'))
entry_tasa_dia.grid(row=8, column=1, padx=5, pady=5, sticky=tk.EW)

ttk.Label(frame_contenido, text="Estatus:").grid(row=9, column=0, padx=5, pady=5, sticky=tk.E)
entry_estatus = ttk.Entry(frame_contenido)
entry_estatus.grid(row=9, column=1, padx=5, pady=5, sticky=tk.EW)
set_entry_text(entry_estatus, "Sin enviar notificación")

# Botón para enviar confirmación
btn_confirmar = ttk.Button(frame_contenido, text="Enviar Confirmación", command=enviar_confirmacion)
btn_confirmar.grid(row=10, column=0, columnspan=2, padx=5, pady=20)

# combo_solicitante.set("Angely Ramírez")
# combo_empresa.set("Tex C.A")
# combo_forma_pago.set("Efectivo Bs")
# combo_banco.set("BANESCO BS (2418)")
# entry_monto_bs.insert(0, "500")

# Iniciar el bot en un hilo separado
threading.Thread(target=iniciar_bot, daemon=True).start()

# Ejecutar la ventana principal
root.mainloop()
